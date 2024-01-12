#!/usr/bin/python3
# Writtin By Thaer Maddah
import sys
import time
import zipfile
import xml.etree.ElementTree as ET
import pandas as pd
import re
sys.path.insert(1, '../')
import write_grades as wr
import browse_files as bf
#import textract
import math


from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
#import zipfile
#import xml.etree.ElementTree as ET
import re
#from docx.shared import Inches
from PIL import Image
from io import BytesIO

data = []

# Check if there is a cover page
def has_cover_page(doc):
    # Iterate through the sections of the document
    for section in doc.sections:
        # Check if the section has a header
        if section.header:
            header = section.header
            # Iterate through the paragraphs in the header
            for paragraph in header.paragraphs:
                # Check if the paragraph contains text
                if paragraph.text.strip():
                    return True  # If there is text in the header, consider it as a cover page
    return False


# Search for word in text
def has_string(doc):
    # Code optimized by mixtral
    # Ever condition must be in new paragraph 
    return sum(1 for p in doc.paragraphs if "Thaer" in p.text or 
        "Artificial Intelligence" in p.text or
    "ثائر" in p.text or 
        "الذكاء الاصطناعي" in p.text)

    #for paragraph in doc.paragraphs:
    #    # Check if the paragraph contains a table of contents field
    #    if "ثائر" or "Thaer" in paragraph.text:
    #        info += 1
    #    if "Artificial" in paragraph.text:
    #        info += 1
    #    return info


    #return False


# This is the true function that's detect table of contents
def has_table_of_contents(doc):
    body_element = doc._body._body
    #print(body_element.xml)
    # Search for table of contents (TOC)
    if "TOC" in body_element.xml:
        return True
    return False

# find table of contents using regular expression revised by claude 2.0
def has_table_of_contents1(doc):
    """Check if a document contains a table of contents"""
    body_text = doc._body._body.text    # we can replace _body.text with _body.xml
    
    # Use a regular expression to search for 'Table of Contents' or 'Contents'
    toc_regex = re.compile(r'(Table of Contents|Contents)', re.IGNORECASE)
    if toc_regex.search(body_text):
        return True
    
    return False

# Font section
def check_font_size(doc):
    body_element = doc._body._body
    #print(body_element.xml)
    return "w:val=\"40\"" in body_element.xml and "w:val=\"28\"" in body_element.xml

def is_bold(doc):
    body_element = doc._body._body
    return "" in body_element.xml

########################## bold section #########################
def is_font_bold(run):
    """
    Check if the font in a run is bold.
    """
    return run.bold

def check_bold_in_word_file(doc):
    """
    Check if the font is bold in a Word file.
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if is_font_bold(run):
                #print(f"Font in paragraph '{paragraph.text}' is bold.")
                return True
    return False
############################################################################

def check_complex_style(doc):

    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph has a specific font complex script style
        #print(paragraph.style.font.name)
        if paragraph.style.font.name == "Times New Roman" or paragraph.style.font.name == "Arial":
            #print(f"The following paragraph has a complex script font style: {paragraph.text}")
            return True
    return False

def check_line_spacing(doc):
    for paragraph in doc.paragraphs:
        # Check line spacing for each paragraph
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing is not None and line_spacing == 1.5:
            #print(f"Paragraph: '{paragraph.text}' - Line Spacing: {line_spacing}")
            return True
    return False


def is_font_justified(paragraph):
    # Check if the paragraph alignment is justified
    return paragraph.alignment == 3  # 3 corresponds to justified alignment

def check_justified_font(doc):
    for paragraph in doc.paragraphs:
        if is_font_justified(paragraph):
            print(f"Paragraph '{paragraph.text}' \nhas justified font.")
            return True
    return False


# Check font color
def is_red_color(color):
    # return red color
    return color == RGBColor(255, 0, 0)

def has_gradient_font_color(color):
    # Check if the color has variations in RGB components
    return len(set(color) - {0, 255}) > 1

def is_red_variant(color):
    # Define a range of RGB values corresponding to red variants
    red_range = range(180, 256)     # Adjust the range as needed
    green_range = range(0, 75)      # Adjust the range as needed
    blue_range = range(0, 75)       # Adjust the range as needed

    # Check if the color is a red variant
    return color[0] in red_range and color[1] in green_range and color[2] in blue_range

def check_font_color(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font_color = run.font.color.rgb if run.font.color else None

            if font_color is not None and is_red_variant(font_color):
                print(f"Red variant font color found in paragraph '{paragraph.text}'")
                return True
    return False


def has_images(doc):
    for rel in doc.part.rels:
        if 'image' in doc.part.rels[rel].target_ref:
            return True
    return False


def get_image_dimensions(image_bytes):
    with Image.open(BytesIO(image_bytes)) as img:
        return img.size


def find_images_with_dimensions(doc, target_dimensions_cm):
    matching_images = []

    # Convert target dimensions from cm to pixels
    target_dimensions_px = (int(target_dimensions_cm[0] / 2.54 * 96), int(target_dimensions_cm[1] / 2.54 * 96))
    #print(target_dimensions_px)

    for rel_id in doc.part.rels:
        rel = doc.part.rels[rel_id]
        if 'image' in rel.target_ref:
            image_part = rel.target_part
            image_dimensions = get_image_dimensions(image_part.blob)
            print(image_dimensions, rel_id)
            
            if image_dimensions == target_dimensions_px:
                matching_images.append(rel_id)

    return matching_images

# References

def has_hyperlinks(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.element.findall('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                return True  # Document has at least one hyperlink

    return False  # No hyperlinks found in the document

def print_hyperlinks(doc):
    hyperlink_count = 0

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            hyperlink = run.element.find('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if hyperlink is not None:
                hyperlink_count += 1
                #url = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                #print(f"Hyperlink {hyperlink_count}: {url}")

    print(f"Total Hyperlinks Found: {hyperlink_count}")



#from docx.opc.constants import RELATIONSHIP_TYPE as RT
#def links(rels):
#    for rel in rels.items():
#        if rels[rel].reltype == RT.HYPERLINK:
#            yield rels[rel]._target


def hyperlinks(doc):
    links = []
    for paragraph in doc.paragraphs:
        if not paragraph.hyperlinks:
            continue
        for hyperlink in paragraph.hyperlinks:
            links.append(hyperlink.address)
    return links


def is_a4_word_document(doc):
    section = doc.sections[0]  # Assuming the first section defines the paper size
    page_width, page_height = section.page_width, section.page_height
    print(page_width, page_height)
    return page_width == 7560310  and page_height == 10692130    # A4 dimensions in twips

def check_margins(doc):
    section = doc.sections[0]  # Assuming the first section defines the margins
    left_margin = section.left_margin
    right_margin = section.right_margin
    top_margin = section.top_margin
    bottom_margin = section.bottom_margin
    print(top_margin, bottom_margin)
    return (
        left_margin == right_margin == 720090  and  # 1440 twips per inch
        top_margin == bottom_margin == 720090 
    )

# Define a function to check watermark
def has_watermark(doc):
    # Loop through the header parts
    for header_part in doc.part.package.parts:
        # Check if the part is a header
        if header_part.partname.find("/header") > 0:
            # Get the xml content of the header
            xml = header_part.blob.decode()
            # Check if the xml contains the watermark tag
            #if "mso-position-horizontal" in xml:
            if "PowerPlusWaterMarkObject" in xml:
                # Return True if watermark is found
                return True
    # Return False if no watermark is found
    return False

def page_has_color(doc):
    # Get the document element
    doc_element = doc.element
    
    # Find the background element
    background = doc_element.find("w:background", doc_element.nsmap)
    
    # Check if the background element exists
    if background is not None:
        # Get the page color
        page_color = background.get(qn("w:color"))
        # Print the page color in hexadecimal format
        print(f"Page color: #{page_color}")
        return True
    else:
        # Print the default page color (white)
        print("Page color: #FFFFFF")
        return False



def reviseDocuments(doc):

    data = []
    # Specify the path to your Word document
    doc_path = '../test/test.docx'

    # Load the Word document
    doc = Document(doc_path)
    print(f"File: {doc_path}")

    # Check if the document has a cover page
    if has_cover_page(doc):
        print("The document has a cover page.")
        data.append(5)
    else:
        print("No cover page found in the document.")
        data.append(0)

    if has_string(doc):
        print(f"The document has string {has_string(doc)}")
        data.append(has_string(doc))
    else:
        print(f"The document has string {has_string(doc)}")
        data.append(0)

    # Check if the document has a table of contents
    if has_table_of_contents(doc):
        print("The document has a table of contents.")
        data.append(5)
    else:
        print("No table of contents found in the document.")
        data.append(0)

    if check_complex_style(doc):
        print("Font name Arial or New Times Roman exists")
        data.append(1)
    else:
        print("Font did not match")
        data.append(0)


    if check_font_size(doc):
        print("Font size: 20 and 14")
        data.append(2)
    else:
        print("Font size it's not correct")
        data.append(0)

    if check_bold_in_word_file(doc):
        print("The text is bold.")
        data.append(1)
    else:
        print("Found no bold")
        data.append(0)

    if check_font_color(doc):
        print("Font color is red")
        data.append(2)
    else:
        print("Font color is not red")
        data.append(1) # add 1 for balck font

    if check_line_spacing(doc):
        print("Line spacing is 1.5")
        data.append(1)
    else:
        print("Line spacing is not correct")
        data.append(0)

    if check_justified_font(doc):
        print('The paragraph is justified')
        data.append(4)
    else:
        print('The paragraph is not justified')
        data.append(0)

    if has_images(doc):
        print("The document contains images.")
    else:
        print("No images found in the document")

    target_dimensions_cm = (6, 6)  # Assuming default DPI is 96, adjust if needed
    try:
        matching_images = find_images_with_dimensions(doc, target_dimensions_cm)
        print(matching_images)

        if matching_images:
            print(f"Images with dimensions {target_dimensions_cm} cm found:")
            for i, rel_id in enumerate(matching_images, start=1):
                print(f"Image {i}: Relationship ID = {rel_id}")
        else:
            print(f"No images with dimensions {target_dimensions_cm} cm found in the document.")
    except Exception as e:
        print(f"An error occurred: {e}")


    all_hyperlinks = hyperlinks(doc)
    if all_hyperlinks: 
        print(f"The document has links\n {all_hyperlinks}")
        data.append(2)
    else:
        print('No hyperlinks found in the document')
        data.append(0)

    if is_a4_word_document(doc):
        print("The Word document has A4 paper size.")
        data.append(2)
    else:
        print("The Word document does not have A4 paper size.")
        data.append(0)

    if check_margins(doc):
        print("The Word document has 2 cm margins on all sides.")
        data.append(2)
    else:
        print("The Word document does not have 2 cm margins on all sides.")
        data.append(0)
  
    if has_watermark(doc):
        print("The Word document contains a watermark.")
        data.append(2)
    else:
        print("The Word document does not contain a watermark.") 
        data.append(0)
    
    if page_has_color(doc):
        data.append(2)
    else:
        data.append(0)

    print(data)
    degree = sum(data[1:len(data)])
    # put data between bracktes to retern data in a single list element [item1,item2,..]
    wr.writeDocGrades([data]) 
    print('Final degree is:', degree)
    del data[:]


def main():
    counter = 0 
    folder = '../Assign/c18'
    #folder = 'test/'
    ext = '.docx'
    trim_txt = '../code/Assign/'
    files = []
    files, dirs = bf.browse(ext, folder)
    sep = '='
    start = time.time()
    
    for file, dir in zip(files, dirs):
        path = bf.getFile(file, dir)
        print(path)
        #print('The file path:', path)
        doc = zipfile.ZipFile(path).read('word/document.xml')
        root = ET.fromstring(doc)
        #xmlfile = ET.parse(path)
        #root = xmlfile.getroot()
        
        f = open(dir +'/' + 'doc.xml', 'wb')
        f.write(doc)
        # Microsoft's XML makes heavy use of XML namespaces; thus, we'll need to reference that in our code
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        body = root.find('w:body', ns)  # find the XML "body" tag
        p_sections = body.findall('w:p', ns)  # under the body tag, find all the paragraph sections
    
        # add student file
        data.append(path.strip(trim_txt))
        reviseDocuments(doc)
        counter += 1
        print(counter, 'Assignments has been revised!')
        print(sep * 120)
        #time.sleep(0.5)
    end = time.time()
    print(f"Total time: {round(end - start)} seconds")



if __name__ == "__main__":
    sys.exit(main())
