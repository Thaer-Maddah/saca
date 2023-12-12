from docx import Document

def read_word_file(file_path):
    doc = Document(file_path)


    for paragraph in doc.paragraphs:
        font_name = None
        font_size = None


        # Check if the paragraph has runs (formatting)
        if paragraph.runs:
            # Get font name from the first run
            #font_name = paragraph.runs[0].font.name
            font_name = paragraph.runs[0].font.name
            # Get font size from the first run
            font_size = paragraph.runs[0].font.size


        print(f"Paragraph: {paragraph.text}")
        print(f"Font Name: {font_name}")
        print(f"Font Size: {font_size}")
        print("\n")
def check_complex_style1(file_path):
    doc = Document(file_path)

    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph has a specific font complex script style
        print(paragraph.style.font.name)
        if paragraph.style.font.name == "Times New Roman":
            print(f"The following paragraph has a complex script font style: {paragraph.text}")


def check_complex_style2(file_path):
    doc = Document(file_path)

    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if any run within the paragraph uses a complex script font
        for run in paragraph.runs:
            # to fix complex script problem I downloaded python-docx source code from github
            # https://github.com/python-openxml/python-docx.git
            # and added to src/oxml/text/font.py  into class CT_Fonts element reads complex script font
            if run._element.rPr.rFonts.cs:
                print(f"The following paragraph contains a complex script font: {paragraph.text}")
                print(run.element.rPr.rFonts.cs)
                #print(run._element.rPr.cs)
                #break  # No need to continue checking runs in this paragraph if one was found


def check_complex_style3(file_path):
    doc = Document(file_path)
    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if any run within the paragraph contains non-ASCII characters
        for run in paragraph.runs:
            text = run.text
            if any(ord(char) > 127 for char in text):
                print(run.font.name_ascii)
                print(f"The following paragraph contains complex script: {paragraph.text}")
                break  # No need to continue checking runs in this paragraph if one was found


def check_urls(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        # Check if any run within the paragraph contains non-ASCII characters
        for run in paragraph.runs:
            text = run.text.hyperlinks
            if text :
                print("there are liks")
            


if __name__ == "__main__":
    # Replace 'your_word_file.docx' with the actual path to your Word file
    word_file_path = "../test/test1.docx"
    #read_word_file(word_file_path)
    check_complex_style2(word_file_path)
    #check_urls(word_file_path)
